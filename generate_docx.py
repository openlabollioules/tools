"""
title: Generate DOCX Document
author: openlab
version: 0.1
license: MIT
description: Génère un fichier DOCX via un LLM (Ollama) et renvoie un lien de téléchargement
"""

import os, uuid
from typing import Optional, Callable, Any
from pathlib import Path
from fastapi import UploadFile, Request
import re
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from open_webui.routers.files import upload_file
from open_webui.models.users import Users
from open_webui.storage.provider import Storage
from open_webui.models.files import Files, FileForm

class EventEmitter:
    def __init__(self, event_emitter: Callable[[dict], Any] = None):
        self.event_emitter = event_emitter

    async def emit(self, description="Unknown State", status="in_progress", done=False):
        if self.event_emitter:
            await self.event_emitter(
                {
                    "type": "status",
                    "data": {
                        "status": status,
                        "description": description,
                        "done": done,
                    },
                }
            )

class HelpFunctions:
    def __init__(self):
        self.styles = {
            "title": "Title",
            "heading1": "Heading 1",
            "heading2": "Heading 2",
            "heading3": "Heading 3",
            "normal": "Normal",
            "subtitle": "Subtitle",
            "toc": "TOC Heading",
        }
        self.fonts = {
            "main": "Calibri",
            "heading": "Arial",
            "title": "Arial",
        }

    def remove_tags_no_keep(self, text: str, start: str, end: str) -> str:
        """
        Remove all text between two tags (`start` and `end`), tags included.

        Parameters
        ----------
        text : str
            text to remove tags from
        start : str
            starting tag
        end : str
            ending tag

        Returns
        -------
        str
            Text with tags removed
        """
        return re.sub(r'{}.*?{}'.format(start, end), '', text, flags=re.DOTALL).strip()

    def upload_file(self, file: UploadFile, user_id: str):
        """
        upload a file to the openwebui data base without the API (API doesn't work with the tools in version 0.6.5)
        ARGS:
            file: the file to upload
            user_id: the id of the user
        RETURNS:
            the file item
        """
        filename = file.filename
        id = str(uuid.uuid4())
        name = filename
        filename = f"{id}_{filename}"
        contents, file_path = Storage.upload_file(file.file, filename)

        file_item = Files.insert_new_file(
            user_id,
            FileForm(
                **{
                    "id": id,
                    "filename": name,
                    "path": file_path,
                    "meta": {
                        "name": name,
                        "content_type": file.content_type,
                        "size": len(contents),
                        "data": {"generated_by": "upload_file"},
                    },
                }
            ),
        )
        print(f"[DEBUG] File item: {file_item}")

        return file_item

    def setup_document_styles(self, doc: Document) -> None:
        """
        Configure document-wide styles for professional formatting
        
        Args:
            doc (Document): The Word document object.
        """
        # Default paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.fonts["main"]
        font.size = Pt(11)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)
        
        # Heading styles
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = self.fonts["heading"]
            font.size = Pt(16 - (i-1)*2)  # 16pt, 14pt, 12pt
            font.bold = True
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            
        # Title style
        style = doc.styles['Title']
        font = style.font
        font.name = self.fonts["title"]
        font.size = Pt(24)
        font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_after = Pt(24)

        # Subtitle style
        style = doc.styles['Subtitle']
        font = style.font
        font.name = self.fonts["title"]
        font.size = Pt(18)
        font.italic = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_after = Pt(36)

    def add_cover_page(self, doc: Document, title: str, subtitle: Optional[str] = None, 
                      author: Optional[str] = None, date: Optional[str] = None, 
                      logo_path: Optional[str] = None) -> None:
        """
        Adds a professional cover page to the document.

        Args:
            doc (Document): The Word document object.
            title (str): The main title of the document.
            subtitle (str, optional): The subtitle or theme.
            author (str, optional): The author or company name.
            date (str, optional): The document date.
            logo_path (str, optional): Path to a logo image.
        """
        # Add page break if document already has content
        if len(doc.paragraphs) > 0 and doc.paragraphs[0].text.strip():
            doc.add_page_break()
        
        # Add a blank paragraph for spacing at top
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add title
        title_para = doc.add_paragraph(title, style='Title')
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add subtitle if provided
        if subtitle:
            subtitle_para = doc.add_paragraph(subtitle, style='Subtitle')
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add logo if path is provided
        if logo_path and os.path.exists(logo_path):
            try:
                doc.add_picture(logo_path, width=Inches(2))
                # Center the picture
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                last_paragraph.space_after = Pt(24)
            except Exception as e:
                print(f"Error adding logo: {e}")
        
        # Add author information
        if author:
            author_para = doc.add_paragraph(f"Auteur: {author}")
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            author_para.space_after = Pt(12)
        
        # Add date
        if date:
            date_para = doc.add_paragraph(date)
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add page break to start the main content on a new page
        doc.add_page_break()

    def add_table_of_contents(self, doc: Document, title: str = "Table des matières") -> None:
        """
        Adds a table of contents to the document.
        
        Args:
            doc (Document): The Word document object.
            title (str, optional): Title for the TOC. Defaults to "Table des matières".
        """
        # Add heading for TOC
        toc_heading = doc.add_paragraph(title, style=self.styles["heading1"])
        toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add TOC field
        para = doc.add_paragraph()
        run = para.add_run()
        
        # Add the TOC field code - corrected implementation
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Includes headings 1-3, with hyperlinks
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # Append all elements properly
        run._element.append(fldChar)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        
        # Add page break after TOC
        doc.add_page_break()

    def add_title(self, doc: Document, title: str = "Title") -> None:
        """
        Adds a title to the document.

        Args:
            doc (Document): The Word document object.
            title (str, optional): The title text. Defaults to "Title".

        Returns:
            None

        Raises:
            ValueError: If the provided title is empty.
        """
        # Check if the title is not empty
        if not title.strip():
            raise ValueError("Document title cannot be empty.")
            
        # Add title
        title_paragraph = doc.add_paragraph(title, style=self.styles["title"])
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_heading(self, doc: Document, heading: str, level: int = 1) -> None:
        """
        Adds a heading to the document.

        Args:
            doc (Document): The Word document object.
            heading (str): The heading text.
            level (int, optional): The heading level (1-3). Defaults to 1.

        Returns:
            None

        Raises:
            ValueError: If the heading is empty or level is invalid.
        """
        # Check if the heading is not empty
        if not heading.strip():
            raise ValueError("Heading cannot be empty.")
            
        # Validate heading level
        if level < 1 or level > 3:
            raise ValueError("Heading level must be between 1 and 3.")
            
        # Map level to style
        style = f"heading{level}"
        
        # Add heading
        doc.add_paragraph(heading, style=self.styles[style])

    def add_paragraph_text(self, doc: Document, content: str) -> None:
        """
        Adds a paragraph of text to the document, handling bullet points.

        Args:
            doc (Document): The Word document object.
            content (str): The paragraph content.

        Returns:
            None
        """
        if not content.strip():
            return
            
        lines = content.split("\n")
        
        for line in lines:
            level = 0
            # Handle indentation
            while line.startswith('    '):
                line = line[4:]
                level += 1
                
            # Handle bullet points
            if line.startswith('* ') or line.startswith('• '):
                line = line[2:]
                p = doc.add_paragraph(line, style='List Bullet')
                p.paragraph_format.left_indent = Pt(level * 18)
            else:
                p = doc.add_paragraph(line)
                if level > 0:
                    p.paragraph_format.left_indent = Pt(level * 18)
                # Apply justified alignment and proper line spacing for regular paragraphs
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_section_header(self, doc: Document, title: str) -> None:
        """
        Adds a formatted section header (like Introduction, Conclusion, etc.)
        
        Args:
            doc (Document): The Word document object.
            title (str): The section header text.
        """
        # Add heading with proper style
        header = doc.add_paragraph(title, style=self.styles["heading1"])
        # Add a subtle horizontal line under the section header
        # La méthode bottom_border n'existe pas directement, utilisons une approche alternative
        pPr = header._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Taille de la bordure (~1pt)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')  # Couleur noire
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        header.paragraph_format.space_after = Pt(12)

    def add_bibliography(self, doc: Document, references: list) -> None:
        """
        Adds a bibliography section with formatted references
        
        Args:
            doc (Document): The Word document object.
            references (list): List of reference strings
        """
        if not references:
            return
            
        # Add section header
        self.add_section_header(doc, "Bibliographie / Références")
        
        # Add each reference as a properly formatted paragraph
        for ref in references:
            p = doc.add_paragraph(ref)
            p.paragraph_format.first_line_indent = Pt(-36)  # Hanging indent
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.space_after = Pt(6)

# --- Tools ---
class Tools:
    def __init__(self):
        self.FILES_DIR = "./tmp"
        self.API_BASE_URL = "http://localhost:8080/api/v1/files/"
        self.template_path = "./templates/template.docx"
        os.makedirs(self.FILES_DIR, exist_ok=True)
        self.help_functions = HelpFunctions()
        self.event_emitter = EventEmitter()
    
    async def generate_docx_from_json(self, json_data: dict, __request__: Request, __event_emitter__: Callable[[dict], Any] = None, __user__=None):
        """
        Generate a Word document from a JSON file.
        
        Args:
            json_data: The JSON data to generate the document from.
            __user__: The user to upload the file to.
        Returns:
            str: The download URL of the uploaded file.
        
        Example:
        Here is an example of the JSON data for a complete document:
            
        ```json
        {
            "titre": "Titre du document",
            "sous_titre": "Sous-titre explicatif",
            "auteur": "Nom de l'auteur ou entreprise",
            "date": "Octobre 2023",
            "logo_path": "chemin/vers/logo.png",
            "inclure_table_matieres": true,
            "sections": [
                {
                    "type": "page_garde",
                    "titre": "Titre principal",
                    "sous_titre": "Sous-titre ou description",
                    "auteur": "Nom ou société",
                    "date": "Date de publication"
                },
                {
                    "type": "introduction",
                    "contenu": "Texte d'introduction...\n* Point important\n* Autre point important"
                },
                {
                    "type": "heading",
                    "titre": "Première partie",
                    "niveau": 1
                },
                {
                    "type": "contenu",
                    "contenu": "Contenu de la première partie...\n* Bullet list\n    * Bullet list niveau 2"
                },
                {
                    "type": "heading",
                    "titre": "Sous-section",
                    "niveau": 2
                },
                {
                    "type": "conclusion",
                    "contenu": "Texte de conclusion..."
                },
                {
                    "type": "bibliographie",
                    "references": [
                        "Auteur, A. (2023). Titre de l'ouvrage. Éditeur.",
                        "Auteur, B. (2022). Titre de l'article. Journal, 10(2), 45-67."
                    ]
                }
            ]
        }
        ```
        """
        emitter = EventEmitter(__event_emitter__)
        print("[DEBUG] json_data", json_data)
        topic = json_data.get('titre')
        print("[DEBUG] topic", topic)
        await emitter.emit(f"Initiating DOCX generation for topic: {topic}")
        
        # Create document
        try:
            doc = Document(self.template_path)
        except:
            doc = Document()  # Create a new document if template doesn't exist
        print("[DEBUG] doc created")
        
        # Set up professional document styles
        self.help_functions.setup_document_styles(doc)
        print("[DEBUG] document styles configured")
        
        # Process document structure in order
        try:
            await emitter.emit("Creating document structure")
            
            # 1. Add cover page if data is provided
            cover_data = next((s for s in json_data.get('sections', []) if s.get('type') == 'page_garde'), None)
            if cover_data:
                self.help_functions.add_cover_page(
                    doc,
                    title=cover_data.get('titre', json_data.get('titre')),
                    subtitle=cover_data.get('sous_titre', json_data.get('sous_titre')),
                    author=cover_data.get('auteur', json_data.get('auteur')),
                    date=cover_data.get('date', json_data.get('date')),
                    logo_path=json_data.get('logo_path')
                )
                print("[DEBUG] cover page added")
            
            # 2. Add table of contents if requested
            if json_data.get('inclure_table_matieres', False):
                self.help_functions.add_table_of_contents(doc)
                print("[DEBUG] table of contents added")
            
            # 3. Process each section in order
            for section in json_data.get('sections', []):
                section_type = section.get('type')
                print(f"[DEBUG] Processing section type: {section_type}")
                
                if section_type == "page_garde":
                    # Already handled above
                    continue
                
                elif section_type == "introduction":
                    self.help_functions.add_section_header(doc, "Introduction")
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print("[DEBUG] introduction added")
                
                elif section_type == "heading":
                    level = section.get('niveau', 1)
                    self.help_functions.add_heading(doc, heading=section.get('titre'), level=level)
                    print(f"[DEBUG] heading level {level} added: {section.get('titre')}")
                
                elif section_type == "contenu":
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print("[DEBUG] content paragraph added")
                
                elif section_type == "conclusion":
                    self.help_functions.add_section_header(doc, "Conclusion")
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print("[DEBUG] conclusion added")
                
                elif section_type == "bibliographie":
                    self.help_functions.add_bibliography(doc, section.get('references', []))
                    print("[DEBUG] bibliography added")
            
            # 4. Add page numbers to the document (footer)
            sections = doc.sections
            for section in sections:
                footer = section.footer
                paragraph = footer.paragraphs[0]
                paragraph.text = "Page "
                run = paragraph.add_run()
                
                # Correct way to create field codes for page numbers
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE'
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._element.append(fldChar1)
                run._element.append(instrText)
                run._element.append(fldChar2)
                
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            await emitter.emit(
                status="complete",
                description=f"DOCX generation completed",
                done=True,
            )
        except Exception as e:
            print("[DEBUG] Error", e)
            return f"Error: {str(e)}"
        
        # Save document
        if not os.path.exists(self.FILES_DIR):
            os.makedirs(self.FILES_DIR)
        # clean up title for filename
        clean_title = re.sub(r'[^\w\s]', '', json_data.get('titre', 'document'))
        clean_title = clean_title.replace(' ', '_')

        output_path = self.FILES_DIR + '/' + clean_title + '.docx'
        doc.save(output_path)
        print("[DEBUG] output_path", output_path)

        try:
            with open(output_path, 'rb') as f:
                print("[DEBUG] f", f)
                files = UploadFile(file=f, filename=os.path.basename(output_path))
                print("[DEBUG] files", files)
                file_item = await self.upload_file(file=files, user_id=__user__['id'], __request__=__request__, __user__=__user__, __event_emitter__=__event_emitter__)
                print("[DEBUG] file_item", file_item)
                return file_item
        except Exception as e:
            print("[DEBUG] Error", e)
            return "Error"

    async def upload_file(self, file: UploadFile, user_id: str, __request__: Request, __user__: dict, __event_emitter__: Callable[[dict], Any] = None):
        emitter = EventEmitter(__event_emitter__)
        metadata = {"data": {"generated_by": "upload_file"}}
 
        await emitter.emit(f"Getting download link for file: {file.filename}")
        
        # get the user for permissions
        user = Users.get_user_by_id(id=__user__['id'])
        print("[DEBUG] user", user)
        # upload the file to the database
        doc = upload_file(request=__request__, file=file, user=user, file_metadata=metadata, process=False)
        print("[DEBUG] doc", doc)

        # get the download link
        download_link = f"{self.API_BASE_URL}{doc.id}/content"
        print("[DEBUG] download_link", download_link)
        await emitter.emit(
                status="complete",
                description=f"Finished generating the DOCX file",
                done=True
            )
        return (
            f"<source><source_id>{doc.filename}</source_id><source_context>" 
            + str(download_link)
            + "</source_context></source>\n"
        )
