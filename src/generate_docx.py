"""
title: Generate DOCX Document
author: openlab
version: 0.1
description: Génère un fichier DOCX via un LLM (Ollama) et renvoie un lien de téléchargement
"""

import os, uuid
from typing import Optional, Callable, Any
from pathlib import Path
from fastapi import UploadFile, Request
import re
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pydantic import BaseModel, Field
from open_webui.routers.files import upload_file
from open_webui.models.users import Users
from open_webui.storage.provider import Storage
from open_webui.models.files import Files, FileForm

class EventEmitter:
    def __init__(self, event_emitter: Callable[[dict], Any] = None):
        self.event_emitter = event_emitter

    async def emit(self, description="Unknown State", status="in_progress", done=False):
        if self.event_emitter:
            await self.event_emitter(
                {
                    "type": "status",
                    "data": {
                        "status": status,
                        "description": description,
                        "done": done,
                    },
                }
            )

class HelpFunctions:
    def __init__(self):
        # Styles basés sur l'analyse du template templates_new.docx
        self.styles = {
            "title": "Section",                     # Style principal pour les titres
            "subtitle": "Normal",                   # Style pour les sous-titres
            "heading1": "Titre1-Numeroté",          # Titre niveau 1
            "heading2": "Titre2-Numéroté",          # Titre niveau 2
            "heading3": "Titre3-Numéroté",          # Titre niveau 3
            "heading4": "Heading 4",                # Titre niveau 4
            "heading5": "Heading 5",                # Titre niveau 5
            "normal": "Normal",
            "paragraphe_standard": "Paragraphe standard",  # Style pour le contenu principal
            "section": "Section",                   # Style pour les sections
            "caption": "Caption",                   # Style pour les légendes
            "table_of_figures": "table of figures"  # Style pour tables des figures
        }
        self.fonts = {
            "main": "Calibri",
            "heading": "Arial",
            "title": "Arial",
        }

    def remove_tags_no_keep(self, text: str, start: str, end: str) -> str:
        """
        Remove all text between two tags (`start` and `end`), tags included.

        Parameters
        ----------
        text : str
            text to remove tags from
        start : str
            starting tag
        end : str
            ending tag

        Returns
        -------
        str
            Text with tags removed
        """
        return re.sub(r'{}.*?{}'.format(start, end), '', text, flags=re.DOTALL).strip()

    def upload_file(self, file: UploadFile, user_id: str):
        """
        upload a file to the openwebui data base without the API (API doesn't work with the tools in version 0.6.5)
        ARGS:
            file: the file to upload
            user_id: the id of the user
        RETURNS:
            the file item
        """
        filename = file.filename
        id = str(uuid.uuid4())
        name = filename
        filename = f"{id}_{filename}"
        contents, file_path = Storage.upload_file(file.file, filename)

        file_item = Files.insert_new_file(
            user_id,
            FileForm(
                **{
                    "id": id,
                    "filename": name,
                    "path": file_path,
                    "meta": {
                        "name": name,
                        "content_type": file.content_type,
                        "size": len(contents),
                        "data": {"generated_by": "upload_file"},
                    },
                }
            ),
        )
        print(f"[DEBUG] File item: {file_item}")

        return file_item

    def setup_document_styles(self, doc: Document) -> None:
        """
        Configure document-wide styles for professional formatting
        
        Args:
            doc (Document): The Word document object.
        """
        # Default paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.fonts["main"]
        font.size = Pt(11)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)
        
        # Heading styles
        for i in range(1, 4):
            style = doc.styles[f'Heading {i}']
            font = style.font
            font.name = self.fonts["heading"]
            font.size = Pt(16 - (i-1)*2)  # 16pt, 14pt, 12pt
            font.bold = True
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            
        # Title style
        style = doc.styles['Title']
        font = style.font
        font.name = self.fonts["title"]
        font.size = Pt(24)
        font.bold = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_after = Pt(24)

        # Subtitle style
        style = doc.styles['Subtitle']
        font = style.font
        font.name = self.fonts["title"]
        font.size = Pt(18)
        font.italic = True
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.paragraph_format.space_after = Pt(36)

    def add_cover_page(self, doc: Document, title: str, subtitle: Optional[str] = None, 
                      author: Optional[str] = None, date: Optional[str] = None, 
                      logo_path: Optional[str] = None) -> None:
        """
        Adds a professional cover page to the document.

        Args:
            doc (Document): The Word document object.
            title (str): The main title of the document.
            subtitle (str, optional): The subtitle or theme.
            author (str, optional): The author or company name.
            date (str, optional): The document date.
            logo_path (str, optional): Path to a logo image.
        """
        # Add page break if document already has content
        if len(doc.paragraphs) > 0 and doc.paragraphs[0].text.strip():
            doc.add_page_break()
        
        # Add a blank paragraph for spacing at top
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add title (use available style or create manually)
        try:
            title_para = doc.add_paragraph(title, style='Title')
        except:
            title_para = doc.add_paragraph(title)
            run = title_para.runs[0]
            run.bold = True
            run.font.size = Pt(24)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add subtitle if provided
        if subtitle:
            try:
                subtitle_para = doc.add_paragraph(subtitle, style='Subtitle')
            except:
                subtitle_para = doc.add_paragraph(subtitle)
                run = subtitle_para.runs[0]
                run.italic = True
                run.font.size = Pt(18)
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add logo if path is provided
        if logo_path and os.path.exists(logo_path):
            try:
                doc.add_picture(logo_path, width=Inches(2))
                # Center the picture
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                last_paragraph.space_after = Pt(24)
                # Add caption if needed
                # caption = doc.add_paragraph("Figure - Logo", style=self.styles["caption"])
            except Exception as e:
                print(f"Error adding logo: {e}")
        
        # Add author information
        if author:
            author_para = doc.add_paragraph(f"Auteur: {author}")
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            author_para.space_after = Pt(12)
        
        # Add date
        if date:
            date_para = doc.add_paragraph(date)
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add page break to start the main content on a new page
        doc.add_page_break()

    def add_table_of_contents(self, doc: Document, title: str = "Table des matières") -> None:
        """
        Adds a table of contents to the document.
        
        Args:
            doc (Document): The Word document object.
            title (str, optional): Title for the TOC. Defaults to "Table des matières".
        """
        # Add heading for TOC
        toc_heading = doc.add_paragraph(title, style=self.styles["section"])
        toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add TOC field
        para = doc.add_paragraph()
        run = para.add_run()
        
        # Add the TOC field code - corrected implementation
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Includes headings 1-3, with hyperlinks
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # Append all elements properly
        run._element.append(fldChar)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        
        # Add page break after TOC
        doc.add_page_break()

    def add_title(self, doc: Document, title: str = "Title") -> None:
        """
        Adds a title to the document.

        Args:
            doc (Document): The Word document object.
            title (str, optional): The title text. Defaults to "Title".

        Returns:
            None

        Raises:
            ValueError: If the provided title is empty.
        """
        # Check if the title is not empty
        if not title.strip():
            raise ValueError("Document title cannot be empty.")
            
        # Add title
        title_paragraph = doc.add_paragraph(title, style=self.styles["title"])
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_heading(self, doc: Document, heading: str, level: int = 1) -> None:
        """Adds a heading to the document."""
        if not heading.strip():
            raise ValueError("Heading cannot be empty.")
            
        if level < 1 or level > 5:
            raise ValueError("Heading level must be between 1 and 5.")
            
        # Map level to style
        style_mapping = {
            1: "heading1",
            2: "heading2", 
            3: "heading3",
            4: "heading4",
            5: "heading5"
        }
        style_key = style_mapping.get(level, "heading1")
        
        # Use actual style names that exist in Word
        actual_styles = {
            "heading1": "Titre1-Numeroté",   # Titre niveau 1
            "heading2": "Titre2-Numéroté",   # Titre niveau 2
            "heading3": "Titre3-Numéroté",   # Titre niveau 3
            "heading4": "Heading 4",         # Titre niveau 4
            "heading5": "Heading 5"          # Titre niveau 5
        }
        
        try:
            style_name = actual_styles.get(style_key, "Titre1-Numeroté")
            doc.add_paragraph(heading, style=style_name)
        except:
            # Fallback to manual formatting if style doesn't exist
            p = doc.add_paragraph(heading)
            p.style = doc.styles['Normal']
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(16 - (level-1)*2)

    def add_paragraph_text(self, doc: Document, content: str) -> None:
        """Adds a paragraph of text to the document, handling bullet points."""
        if not content.strip():
            return
            
        lines = content.split("\n")
        
        for line in lines:
            level = 0
            # Handle indentation
            while line.startswith('    '):
                line = line[4:]
                level += 1
                
            # Handle bullet points
            if line.startswith('* ') or line.startswith('• '):
                line = line[2:]
                try:
                    p = doc.add_paragraph(line, style='List Bullet')
                except:
                    # Fallback if List Bullet style doesn't exist
                    p = doc.add_paragraph(f"• {line}", style='Normal')
                p.paragraph_format.left_indent = Pt(level * 18)
            else:
                # Use appropriate style for regular text
                try:
                    p = doc.add_paragraph(line, style='Paragraphe standard')
                except:
                    p = doc.add_paragraph(line, style='Normal')
                if level > 0:
                    p.paragraph_format.left_indent = Pt(level * 18)
                # Apply justified alignment and proper line spacing for regular paragraphs
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_section_header(self, doc: Document, title: str) -> None:
        """Adds a formatted section header (like Introduction, Conclusion, etc.)"""
        # Add heading with proper style
        try:
            header = doc.add_paragraph(title, style=self.styles["section"])
        except:
            header = doc.add_paragraph(title)
            run = header.runs[0]
            run.bold = True
            run.font.size = Pt(16)
        
        header.paragraph_format.space_after = Pt(12)

    def add_bibliography(self, doc: Document, references: list) -> None:
        """
        Adds a bibliography section with formatted references
        
        Args:
            doc (Document): The Word document object.
            references (list): List of reference strings
        """
        if not references:
            return
            
        # Add section header
        self.add_section_header(doc, "Bibliographie / Références")
        
        # Add each reference as a properly formatted paragraph
        for ref in references:
            p = doc.add_paragraph(ref)
            p.paragraph_format.first_line_indent = Pt(-36)  # Hanging indent
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.space_after = Pt(6)

# --- Tools ---
class Tools:
    class Valves(BaseModel):
        API_BASE_URL: str = Field(
            default="http://localhost:3000/api/v1/files/", description="url for the API"
        )        
        FILES_DIR: str = Field(
            default="./tmp", description="Path to the folder in which files will be saved"
        )
        base_template_path: str = Field(
            default="./templates/docx/templates_new.docx", description="Path to the folder in which the base template will be saved"
        )
        prefix: str = Field(
            default="CS-IN_", description="Prefix for the file name"
        )
    def __init__(self):
        self.valves = self.Valves()
        self.FILES_DIR = self.valves.FILES_DIR
        self.API_BASE_URL = self.valves.API_BASE_URL
        self.template_path = self.valves.template_path
        self.prefix = self.valves.prefix

        os.makedirs(self.FILES_DIR, exist_ok=True)
        self.help_functions = HelpFunctions()
        self.event_emitter = EventEmitter()
    
    async def generate_docx_from_json(self, json_data: dict, __request__: Request, __event_emitter__: Callable[[dict], Any] = None, __user__=None):
        """
        Generate a Word document from a JSON file.
        
        Args:
            json_data: The JSON data to generate the document from.
            __user__: The user to upload the file to.
        Returns:
            str: The download URL of the uploaded file.
        
        Example:
        Here is an example of the JSON data for a complete document:
            
        ```json
        {
            "titre": "Intelligence Artificielle : Enjeux et Perspectives",
            "sous_titre": "Une analyse complète des technologies d'IA modernes",
            "auteur": "OpenLab Research",
            "date": "02/05/2024",
            "logo_path": null,
            "inclure_table_matieres": true,
            "sections": [
                {
                "type": "page_garde",
                "titre": "Intelligence Artificielle : Enjeux et Perspectives",
                "sous_titre": "Une analyse complète des technologies d'IA modernes",
                "auteur": "OpenLab Research",
                "date": "Décembre 2024"
                },
                {
                "type": "introduction",
                "contenu": "L'intelligence artificielle (IA) représente l'une des révolutions technologiques les plus importantes de notre époque. Cette technologie transforme radicalement notre façon de travailler, de communiquer et d'interagir avec le monde qui nous entoure.\n\nCe document explore les principales dimensions de l'IA :\n* Les fondements techniques et théoriques\n* Les applications actuelles dans différents secteurs\n* Les défis éthiques et sociétaux\n* Les perspectives d'avenir et les tendances émergentes\n\nL'objectif est de fournir une vue d'ensemble accessible et complète de cette technologie révolutionnaire."
                },
                {
                "type": "heading",
                "titre": "Fondements de l'Intelligence Artificielle",
                "niveau": 1
                },
                {
                "type": "contenu",
                "contenu": "L'intelligence artificielle repose sur plusieurs piliers fondamentaux qui permettent aux machines de simuler l'intelligence humaine.\n\nLes principales approches incluent :\n* L'apprentissage automatique (Machine Learning)\n* Les réseaux de neurones artificiels\n* Le traitement du langage naturel (NLP)\n* La vision par ordinateur\n* Les systèmes experts\n\nChacune de ces approches contribue à créer des systèmes capables d'analyser, de comprendre et de prendre des décisions de manière autonome."
                },
                {
                "type": "heading",
                "titre": "Applications Actuelles de l'IA",
                "niveau": 1
                },
                {
                "type": "contenu",
                "contenu": "L'IA est aujourd'hui présente dans de nombreux secteurs d'activité, transformant les processus et créant de nouvelles opportunités.\n\nSecteurs d'application majeurs :\n* Santé : diagnostic médical, découverte de médicaments\n* Transport : véhicules autonomes, optimisation logistique\n* Finance : détection de fraude, trading algorithmique\n* Éducation : personnalisation des apprentissages\n* Industrie : maintenance prédictive, contrôle qualité\n* Commerce : recommandations personnalisées, chatbots\n\nChaque secteur bénéficie d'innovations spécifiques adaptées à ses besoins."
                },
                {
                "type": "conclusion",
                "contenu": "L'intelligence artificielle représente une technologie transformatrice qui redéfinit notre rapport au monde numérique. Ses applications actuelles démontrent déjà son potentiel considérable dans de nombreux domaines.\n\nCependant, son développement doit s'accompagner d'une réflexion éthique et sociétale pour garantir que ses bénéfices profitent à l'ensemble de la société."
                },
                {
                "type": "bibliographie",
                "references": [
                    "Russell, S. & Norvig, P. (2020). Artificial Intelligence: A Modern Approach. Pearson.",
                    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
                    "O'Neil, C. (2016). Weapons of Math Destruction. Crown Books.",
                    "Floridi, L. (2019). The Ethics of Information. Oxford University Press."
                ]
                }
            ]
            }
        ```
        """
        emitter = EventEmitter(__event_emitter__)
        print("[DEBUG] json_data", json_data)
        topic = json_data.get('titre')
        print("[DEBUG] topic", topic)
        await emitter.emit(f"Initiating DOCX generation for topic: {topic}")
        
        # Create document
        try:
            doc = Document(self.template_path)
            print(" Template chargé avec succès")
        except Exception as e:
            doc = Document()  # Create a new document if template doesn't exist
            print(f"  Erreur template, nouveau document créé: {e}")
        
        # Set up professional document styles
        self.help_functions.setup_document_styles(doc)
        print(" Styles configurés")
        
        # Process document structure in order
        try:
            await emitter.emit("Creating document structure")
            print("Création de la structure du document...")
            
            # 1. Add cover page if data is provided
            cover_data = next((s for s in json_data.get('sections', []) if s.get('type') == 'page_garde'), None)
            if cover_data:
                self.help_functions.add_cover_page(
                    doc,
                    title=cover_data.get('titre', json_data.get('titre')),
                    subtitle=cover_data.get('sous_titre', json_data.get('sous_titre')),
                    author=cover_data.get('auteur', json_data.get('auteur')),
                    date=cover_data.get('date', json_data.get('date')),
                    logo_path=json_data.get('logo_path')
                )
                print(" Page de garde ajoutée")
            
            # 2. Add table of contents if requested
            if json_data.get('inclure_table_matieres', False):
                self.help_functions.add_table_of_contents(doc)
                print(" Table des matières ajoutée")
            
            # 3. Process each section in order
            for section in json_data.get('sections', []):
                section_type = section.get('type')
                print(f" Traitement section: {section_type}")
                
                if section_type == "page_garde":
                    # Already handled above
                    continue
                
                elif section_type == "introduction":
                    self.help_functions.add_section_header(doc, "Introduction")
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print(" Introduction ajoutée")
                
                elif section_type == "heading":
                    level = section.get('niveau', 1)
                    self.help_functions.add_heading(doc, heading=section.get('titre'), level=level)
                    print(f" Heading niveau {level} ajouté: {section.get('titre')}")
                
                elif section_type == "contenu":
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print(" Contenu ajouté")
                
                elif section_type == "conclusion":
                    self.help_functions.add_section_header(doc, "Conclusion")
                    self.help_functions.add_paragraph_text(doc, section.get('contenu', ''))
                    print(" Conclusion ajoutée")
                
                elif section_type == "bibliographie":
                    self.help_functions.add_bibliography(doc, section.get('references', []))
                    print(" Bibliographie ajoutée")
            
            # 4. Add page numbers to the document (footer)
            sections = doc.sections
            for section in sections:
                footer = section.footer
                paragraph = footer.paragraphs[0]
                paragraph.text = "Page "
                run = paragraph.add_run()
                
                # Correct way to create field codes for page numbers
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE'
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._element.append(fldChar1)
                run._element.append(instrText)
                run._element.append(fldChar2)
                
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            print("Structure du document terminée")
            
            await emitter.emit(
                status="complete",
                description=f"DOCX generation completed",
                done=True,
            )
        except Exception as e:
            print(f" Erreur lors de la création: {str(e)}")
            import traceback
            traceback.print_exc()
            return f"Error: {str(e)}"
        
        # Save document
        if not os.path.exists(self.FILES_DIR):
            os.makedirs(self.FILES_DIR)
        # clean up title for filename
        clean_title = re.sub(r'[^\w\s]', '', json_data.get('titre', 'document'))
        clean_title = clean_title.replace(' ', '_')
        # add the prefix to the title
        output_path = self.FILES_DIR + '/' + self.prefix + clean_title + '.docx'
        doc.save(output_path)
        print(f"Document sauvegardé: {output_path}")

        try:
            with open(output_path, 'rb') as f:
                print("[DEBUG] f", f)
                files = UploadFile(file=f, filename=os.path.basename(output_path))
                print("[DEBUG] files", files)
                file_item = await self.upload_file(file=files, user_id=__user__['id'], __request__=__request__, __user__=__user__, __event_emitter__=__event_emitter__)
                print("[DEBUG] file_item", file_item)
                return file_item
        except Exception as e:
            print("[DEBUG] Error", e)
            return "Error"

    async def upload_file(self, file: UploadFile, user_id: str, __request__: Request, __user__: dict, __event_emitter__: Callable[[dict], Any] = None):
        emitter = EventEmitter(__event_emitter__)
        metadata = {"data": {"generated_by": "upload_file"}}
 
        await emitter.emit(f"Getting download link for file: {file.filename}")
        
        # get the user for permissions
        user = Users.get_user_by_id(id=__user__['id'])
        print("[DEBUG] user", user)
        # upload the file to the database
        doc = upload_file(request=__request__, file=file, user=user, metadata=metadata, process=False) # process false to not analyse the file
        print("[DEBUG] doc", doc)

        # get the download link
        download_link = f"{self.API_BASE_URL}{doc.id}/content"
        print("[DEBUG] download_link", download_link)
        await emitter.emit(
                status="complete",
                description=f"Finished generating the DOCX file",
                done=True
            )
        return (
            f"<source><source_id>{doc.filename}</source_id><source_context>" 
            + str(download_link)
            + "</source_context></source>\n"
        )