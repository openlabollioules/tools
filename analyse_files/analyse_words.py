#!/usr/bin/env python3
# coding: utf-8

from docx import Document
from lxml import etree
import zipfile
import sys

# -- FONCTIONS UTILES --------------------------------------------------------

def get_outline_level(para):
    """
    Retourne l'outline level (niveau hiérarchique) d'un paragraphe
    (0 pour pas de niveau, 1 pour Heading 1, etc.), ou None si absent.
    """
    pPr = para._p.pPr
    if pPr is not None and pPr.outlineLvl is not None:
        return int(pPr.outlineLvl.val)
    return None

def extract_shapes_from_document_xml(docx_path):
    """
    Ouvre le document .docx en tant que ZIP, lit word/document.xml
    et renvoie la liste des noeuds <pic:spPr> (shape properties).
    """
    ns = {
        'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    }
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    # trouve tous les <pic:spPr> sous n'importe quel <w:drawing>
    return root.findall('.//w:drawing//pic:spPr', namespaces=ns), ns

def parse_shape_properties(spPr, ns):
    """
    À partir d'un élément <pic:spPr>, extrait :
      - type de forme (prstGeom@prst)
      - couleur de remplissage (solidFill/srgbClr@val)
      - couleur de contour (ln//solidFill/srgbClr@val)
    """
    # type de forme
    prst = spPr.find('.//a:prstGeom', namespaces=ns)
    shape_type = prst.get('prst') if prst is not None else None

    # remplissage
    fill = spPr.find('.//a:solidFill/a:srgbClr', namespaces=ns)
    fill_color = fill.get('val') if fill is not None else None

    # contour
    ln = spPr.find('.//a:ln//a:solidFill/a:srgbClr', namespaces=ns)
    line_color = ln.get('val') if ln is not None else None

    return {
        'shape_type':  shape_type,
        'fill_color':  fill_color,
        'line_color':  line_color
    }

# -- SCRIPT PRINCIPAL --------------------------------------------------------

def analyze_docx(path):
    doc = Document(path)

    print(f"\n📄 Analyse du document : {path}\n")

    # 1) Paragraphes : style et niveau
    print("– Paragraphes (texte abrégé, style, outline level) –")
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip().replace('\n',' ')[:40]
        lvl  = get_outline_level(para)
        print(f" {i:3d}. « {text}… »  | Style = {para.style.name!r} | Niveau = {lvl}")

    # 2) Styles de caractère utilisés
    char_styles = {
        run.style.name
        for para in doc.paragraphs
        for run in para.runs
        if run.style is not None
    }
    print("\n– Styles de caractère détectés –")
    for name in sorted(char_styles):
        print(f" • {name}")

    # 3) Inline shapes via python-docx (type brut)
    print("\n– Inline shapes (via python-docx) –")
    for i, shp in enumerate(doc.inline_shapes, 1):
        print(f" {i:3d}. Type brut = {shp.type}")

    # 4) Inline shapes – détails styles (via parsing XML)
    print("\n– Détails shapes (via parsing XML) –")
    spPr_list, ns = extract_shapes_from_document_xml(path)
    for i, spPr in enumerate(spPr_list, 1):
        info = parse_shape_properties(spPr, ns)
        print(
            f" {i:3d}. Forme = {info['shape_type']!r} | "
            f"Remplissage = {info['fill_color'] or 'aucune'} | "
            f"Contour     = {info['line_color'] or 'aucun'}"
        )


if __name__ == "__main__":

    # analyze_docx("./templates/docx/CS-IN_Template-old.docx")
    # analyze_docx("./templates/docx/CS-IN_Template.docx")
    analyze_docx("./templates/templates_new.docx")
